# backend/app/services/extract_input_field_service.py
from __future__ import annotations

import re
import unicodedata
from dataclasses import dataclass, field
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any, Iterator, Optional, Sequence

import pandas as pd
import pdfplumber
from docx import Document as DocxDocument
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph as DocxParagraph

from app.services.extract_xlsx_input_service import UnsupportedWorkbookError
from app.services.extract_xlsx_input_service import build_agent_input as build_workbook_agent_input
from app.services.extract_xlsx_input_service import extract_workbook

WORKBOOK_SUFFIXES = {".xlsx", ".xlsm"}


FIELD_ALIASES = {
    "factory_name": ["nama pabrik", "nama perusahaan", "factory name", "plant name", "company name"],
    "process_type": ["jenis proses", "tipe proses", "jenis alur proses", "process type", "type of process"],
    "layout_description": ["deskripsi pabrik", "gambaran umum pabrik", "factory description",
                           "plant description", "gambaran umum", "deskripsi"],
    "worker_count": ["jumlah pekerja", "jumlah tenaga kerja", "total pekerja", "number of workers",
                     "total workers", "headcount"],
}

MIN_EVIDENCE_SCORE = 2

NON_CV_STEMS = {"readme", "read me", "notes", "catatan", "index", "license", "lisensi",
                "changelog", "manifest", "daftar isi", "petunjuk", "instructions"}

COLUMN_ALIASES = {
    "index": ["no", "nomor", "urut"],
    "stage": ["tahapan proses", "tahap proses", "nama tahap", "nama proses", "process stage",
              "workflow step", "tahap", "stage", "step", "proses"],
    "equipment": ["peralatan", "equipment", "mesin", "machine", "alat"],
    "task": ["deskripsi tugas operator", "operator task description", "deskripsi tugas", "task description",
             "uraian tugas", "tugas operator", "job description", "deskripsi kerja"],
    "automation": ["otomatisasi", "otomasi", "automation", "automated", "is automated"],
    "qc": ["quality control qc", "quality control", "pengendalian mutu", "kendali mutu", "qc"],
    "throughput": ["throughput estimasi", "throughput estimated", "throughput", "kapasitas per jam",
                   "laju produksi", "output per jam"],
    "workers_count": ["jumlah pekerja", "jumlah operator", "jumlah tenaga kerja", "number of workers",
                      "worker count", "headcount"],
    "workers_assigned": ["pekerja ditugaskan", "assigned workers", "assigned worker", "penugasan pekerja",
                         "operator ditugaskan", "nama pekerja"],
    "worker_note": ["keterangan penugasan", "assignment notes", "assignment note", "catatan penugasan",
                    "keterangan"],
    "units": ["jumlah unit", "number of units", "unit count", "banyak unit", "jumlah alat"],
    "capacity_per_unit": ["kapasitas per unit", "capacity per unit", "kapasitas unit"],
    "total_capacity": ["total kapasitas tahap", "total stage capacity", "total kapasitas", "kapasitas total",
                       "total capacity"],
}

FIELD_HEADERS = {
    "stage": "Tahap Proses",
    "equipment": "Peralatan",
    "equipment_inventory": "Peralatan",
    "task": "Deskripsi Tugas Operator",
    "automation": "Otomatisasi",
    "qc": "Quality Control (QC)",
    "throughput": "Throughput (estimasi)",
    "workers_count": "Jumlah Pekerja",
    "workers_assigned": "Pekerja Ditugaskan",
    "worker_note": "Keterangan Penugasan",
    "units": "Jumlah Unit",
    "capacity_per_unit": "Kapasitas per Unit",
    "total_capacity": "Total Kapasitas Tahap",
}

TABLE_GROUPS = [
    (1, ["stage", "equipment", "task", "automation", "qc", "throughput"]),
    (2, ["stage", "workers_count", "workers_assigned", "worker_note"]),
    (3, ["stage", "equipment_inventory", "units", "capacity_per_unit", "total_capacity"]),
]

TABLE_LABELS = {
    1: "TABEL 1 (tahap proses, peralatan, deskripsi tugas operator, otomatisasi, QC)",
    2: "TABEL 2 (tahap proses, jumlah pekerja, pekerja ditugaskan)",
    3: "TABEL 3 (tahap proses, peralatan, jumlah unit)",
}

INVENTORY_FIELDS = {"units", "capacity_per_unit", "total_capacity"}
PROCESS_FIELDS = {"task", "qc", "throughput", "automation"}

SECTION_BREAK = re.compile(
    r"^\s*(tabel|table|lampiran|appendix|catatan|note|notes|sumber|source|disusun|profil|rincian|daftar)\b",
    re.IGNORECASE)

SUMMARY_ROW = re.compile(r"^(total|subtotal|jumlah total|grand total|keterangan|catatan|sumber|note)\b",
                         re.IGNORECASE)

NUMBERED_HEADING = re.compile(r"^\s*(?:tahap|stage|step|langkah)?\s*\d+\s*[\.\)\-:]\s*(\S.*)$", re.IGNORECASE)

BULLET_CHARS = "\u2022\u25cf\u25aa\u25a0\u25e6\u2023\u2043"

CID_PATTERN = re.compile(r"\(cid:\d+\)")

RULED_SETTINGS = {
    "vertical_strategy": "lines",
    "horizontal_strategy": "lines",
    "intersection_tolerance": 5,
    "snap_tolerance": 4,
    "join_tolerance": 4,
}

MAX_STAGE_NAME_LENGTH = 60

AUTOMATION_TRUE = {"\u2713", "\u2714", "\u2611", "\u25cf", "v", "x", "*", "ya", "yes", "true",
                   "otomatis", "automatic", "automated", "ada"}
AUTOMATION_FALSE = {"-", "\u2013", "\u2014", "\u2717", "\u2718", "\u2610", "n", "no", "tidak",
                    "manual", "none", "false", "nihil"}


class UnsupportedDocumentError(ValueError):
    pass


@dataclass
class TextLine:
    text: str
    break_before: bool = False


@dataclass
class RawTable:
    headers: list[str]
    rows: list[list[str]]
    order: int
    width: int


@dataclass
class ExtractedTable:
    index: int
    headers: list[str]
    rows: list[list[str]]

    def to_markdown(self) -> str:
        frame = pd.DataFrame(self.rows, columns=self.headers)
        return frame.to_markdown(index=False)


@dataclass
class ExtractedDocument:
    source_name: str
    text_fields: dict[str, str] = field(default_factory=dict)
    tables: list[ExtractedTable] = field(default_factory=list)
    stages: list[dict[str, Any]] = field(default_factory=list)
    raw_text: str = ""

    def workflow_sequence(self) -> list[str]:
        return [stage["stage"] for stage in self.stages]

    def missing_text_fields(self) -> list[str]:
        return [key for key in FIELD_ALIASES if not self.text_fields.get(key)]

    def table_by_index(self, index: int) -> Optional[ExtractedTable]:
        for table in self.tables:
            if table.index == index:
                return table
        return None

    def missing_tables(self) -> list[int]:
        present = {table.index for table in self.tables}
        return [expected for expected in (1, 2, 3) if expected not in present]

    def is_complete(self) -> bool:
        return not self.missing_text_fields() and not self.missing_tables() and bool(self.stages)


def normalize(value: Optional[str]) -> str:
    if not value:
        return ""
    text = unicodedata.normalize("NFKD", str(value))
    text = "".join(ch for ch in text if not unicodedata.combining(ch))
    text = text.lower().replace("&", " dan ")
    text = re.sub(r"[^a-z0-9]+", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def clean_cell(value: Any) -> str:
    if value is None:
        return ""
    text = CID_PATTERN.sub(" ", str(value)).replace("\n", " ").replace("\u00ad", "")
    text = re.sub(r"[\s\u00a0]+", " ", text).strip()
    trimmed = text.lstrip(BULLET_CHARS + " -*\u2013\u2014").strip()
    return trimmed if trimmed else text


def alias_hit(alias: str, norm: str) -> bool:
    if alias in norm:
        return True
    return len(alias) >= 4 and alias.replace(" ", "") in norm.replace(" ", "")


def match_column(value: Optional[str]) -> Optional[str]:
    norm = normalize(value)
    if not norm:
        return None
    best: Optional[tuple[int, str]] = None
    for name, aliases in COLUMN_ALIASES.items():
        for alias in aliases:
            if alias_hit(alias, norm) and (best is None or len(alias) > best[0]):
                best = (len(alias), name)
    return best[1] if best else None


def similar(left: str, right: str) -> float:
    return SequenceMatcher(None, left, right).ratio()


def looks_like_stage_name(value: str) -> bool:
    text = clean_cell(value)
    if not text or len(text) > MAX_STAGE_NAME_LENGTH:
        return False
    if SUMMARY_ROW.match(text) or SECTION_BREAK.match(text):
        return False
    matched = match_column(text)
    if matched is not None and matched != "stage":
        return False
    return bool(re.search(r"[A-Za-z\u00c0-\u024f]", text))


def split_line_cells(line: dict[str, Any], gap: float = 6.0) -> list[dict[str, Any]]:
    cells: list[dict[str, Any]] = []
    for word in line["words"]:
        if cells and word["x0"] - cells[-1]["x1"] <= gap:
            cells[-1]["text"] += " " + word["text"]
            cells[-1]["x1"] = word["x1"]
        else:
            cells.append({"text": word["text"], "x0": word["x0"], "x1": word["x1"]})
    return cells


def header_score(line: dict[str, Any]) -> tuple[int, list[Optional[str]], list[dict[str, Any]]]:
    cells = split_line_cells(line)
    fields = [match_column(cell["text"]) for cell in cells]
    unique = {name for name in fields if name and name != "index"}
    return len(unique), fields, cells


def is_header_fields(fields: Sequence[Optional[str]], width: int) -> bool:
    unique = {name for name in fields if name and name != "index"}
    if len(unique) >= 3:
        return True
    return len(unique) >= 2 and "stage" in unique and width >= 3


def is_workbook(path: str | Path) -> bool:
    return Path(path).suffix.lower() in WORKBOOK_SUFFIXES


def extract_any(path: str | Path):
    if is_workbook(path):
        return extract_workbook(path)
    return extract_document(path)


def build_any_agent_input(source) -> str:
    if isinstance(source, ExtractedDocument):
        return build_agent_input(source)
    return build_workbook_agent_input(source)


class PageModel:
    def __init__(self, page: Any, index: int) -> None:
        self.index = index
        self.width = float(page.width)
        self.words = page.extract_words(keep_blank_chars=False, use_text_flow=False)
        self.lines = self._cluster_lines(self.words)

    @staticmethod
    def _cluster_lines(words: Sequence[dict[str, Any]], tolerance: float = 3.0) -> list[dict[str, Any]]:
        lines: list[dict[str, Any]] = []
        for word in sorted(words, key=lambda item: ((item["top"] + item["bottom"]) / 2.0, item["x0"])):
            center = (word["top"] + word["bottom"]) / 2.0
            if lines and abs(center - lines[-1]["center"]) <= tolerance:
                lines[-1]["words"].append(word)
                lines[-1]["top"] = min(lines[-1]["top"], word["top"])
                lines[-1]["bottom"] = max(lines[-1]["bottom"], word["bottom"])
            else:
                lines.append({"center": center, "top": word["top"], "bottom": word["bottom"], "words": [word]})
        for line in lines:
            line["words"].sort(key=lambda item: item["x0"])
            line["text"] = " ".join(item["text"] for item in line["words"])
            line["x0"] = line["words"][0]["x0"]
            line["x1"] = line["words"][-1]["x1"]
        return lines

    def regions(self) -> list[list[dict[str, Any]]]:
        gutter = self._find_gutter()
        if gutter is None:
            return [self.lines]
        left = [line for line in self.lines if line["x1"] <= gutter]
        right = [line for line in self.lines if line["x0"] >= gutter]
        straddling = [line for line in self.lines if line["x0"] < gutter < line["x1"]]
        if straddling or not left or not right:
            return [self.lines]
        return [left, right]

    def _find_gutter(self) -> Optional[float]:
        spans = sorted((word["x0"], word["x1"]) for word in self.words)
        if not spans:
            return None
        best: Optional[tuple[float, float]] = None
        cursor = spans[0][1]
        for x0, x1 in spans[1:]:
            center = (cursor + x0) / 2.0
            if x0 - cursor > 14 and 0.2 * self.width < center < 0.8 * self.width:
                if best is None or x0 - cursor > best[0]:
                    best = (x0 - cursor, center)
            cursor = max(cursor, x1)
        if best is None:
            return None
        left = [line for line in self.lines if line["x1"] <= best[1]]
        right = [line for line in self.lines if line["x0"] >= best[1]]
        if len(left) < 4 or len(right) < 4:
            return None
        return best[1]


def whitespace_gaps(words: Sequence[dict[str, Any]], minimum: float = 5.0) -> list[float]:
    spans = sorted((word["x0"], word["x1"]) for word in words)
    if not spans:
        return []
    gaps = []
    cursor = spans[0][1]
    for x0, x1 in spans[1:]:
        if x0 - cursor >= minimum:
            gaps.append((cursor + x0) / 2.0)
        cursor = max(cursor, x1)
    return gaps


def refine_bounds(header_cells: Sequence[dict[str, Any]], body: Sequence[dict[str, Any]],
                  page_width: float) -> list[float]:
    seeds = [(header_cells[i]["x1"] + header_cells[i + 1]["x0"]) / 2.0
             for i in range(len(header_cells) - 1)]
    words = [word for line in body for word in line["words"]]
    words += [{"x0": cell["x0"], "x1": cell["x1"]} for cell in header_cells]
    gaps = whitespace_gaps(words)
    if len(gaps) == len(seeds):
        chosen = gaps
    else:
        chosen = []
        floor = -1e6
        for seed in seeds:
            candidates = [gap for gap in gaps if gap > floor and abs(gap - seed) <= 45]
            value = min(candidates, key=lambda gap: abs(gap - seed)) if candidates else max(seed, floor + 1)
            chosen.append(value)
            floor = value
    return [-1e6] + chosen + [page_width + 1e6]





def row_break_threshold(body: Sequence[dict[str, Any]]) -> Optional[float]:
    gaps = [body[i]["top"] - body[i - 1]["bottom"] for i in range(1, len(body))]
    gaps = [gap for gap in gaps if gap > -2]
    if len(gaps) < 2:
        return None
    low, high = min(gaps), max(gaps)
    if high - low <= 2.0:
        return None
    return (low + high) / 2.0


def assign_rows(body: Sequence[dict[str, Any]], bounds: Sequence[float], width: int,
                anchor: int) -> list[list[str]]:
    threshold = row_break_threshold(body)
    rows: list[list[str]] = []
    for position, line in enumerate(body):
        cells = [""] * width
        for word in line["words"]:
            center = (word["x0"] + word["x1"]) / 2.0
            slot = width - 1
            for column in range(width):
                if bounds[column] <= center < bounds[column + 1]:
                    slot = column
                    break
            cells[slot] = (cells[slot] + " " + word["text"]).strip()
        gap = line["top"] - body[position - 1]["bottom"] if position else None
        anchored = bool(cells[anchor].strip()) if anchor < width else bool(cells[0].strip())
        starts_row = not rows or (anchored and (threshold is None or (gap is not None and gap > threshold)))
        if starts_row:
            rows.append(cells)
        else:
            for column, value in enumerate(cells):
                if value:
                    rows[-1][column] = (rows[-1][column] + " " + value).strip()
    return rows


def table_end(region: Sequence[dict[str, Any]], start: int) -> int:
    for index in range(start, len(region)):
        line = region[index]
        text = clean_cell(line["text"])
        first = clean_cell(split_line_cells(line)[0]["text"])
        if SUMMARY_ROW.match(first):
            return index
        if SECTION_BREAK.match(text) and not looks_like_stage_name(first):
            return index
        score, fields, cells = header_score(line)
        if index > start and is_header_fields(fields, len(cells)):
            return index
    return len(region)


def extend_header(region: Sequence[dict[str, Any]], index: int, fields: list[Optional[str]],
                  cells: list[dict[str, Any]]) -> tuple[list[Optional[str]], list[dict[str, Any]], int]:
    merged = [dict(cell) for cell in cells]
    consumed = 1
    for offset in (1, 2):
        position = index + offset
        if position >= len(region):
            break
        candidate = region[position]
        if candidate["top"] - region[position - 1]["bottom"] > 6:
            break
        trial = [dict(cell) for cell in merged]
        for word in candidate["words"]:
            center = (word["x0"] + word["x1"]) / 2.0
            target = None
            for slot in trial:
                if slot["x0"] - 3 <= center <= slot["x1"] + 3:
                    target = slot
                    break
            if target is None:
                target = min(trial, key=lambda slot: min(abs(center - slot["x0"]), abs(center - slot["x1"])))
            target["text"] += " " + word["text"]
            target["x0"] = min(target["x0"], word["x0"])
            target["x1"] = max(target["x1"], word["x1"])
        trial_fields = [match_column(slot["text"]) for slot in trial]
        current = {name for name in fields if name and name != "index"}
        upgraded = {name for name in trial_fields if name and name != "index"}
        if len(upgraded) < len(current):
            break
        merged, fields, consumed = trial, trial_fields, consumed + 1
    return fields, merged, consumed


def ruled_tables(page: Any) -> list[tuple[list[list[str]], tuple[float, float]]]:
    output = []
    for table in page.find_tables(RULED_SETTINGS):
        rows = [[clean_cell(cell) for cell in row] for row in table.extract()]
        rows = [row for row in rows if any(row)]
        if rows and len(rows[0]) >= 2:
            output.append((rows, (table.bbox[1], table.bbox[3])))
    return output


def geometric_tables(model: PageModel) -> list[tuple[list[list[str]], tuple[float, float]]]:
    output = []
    for region in model.regions():
        index = 0
        while index < len(region):
            score, fields, cells = header_score(region[index])
            if is_header_fields(fields, len(cells)) and len(cells) >= 3:
                fields, header_cells, consumed = extend_header(region, index, fields, cells)
                end = table_end(region, index + consumed)
                body = region[index + consumed:end]
                bounds = refine_bounds(header_cells, body, model.width)
                anchor = fields.index("stage") if "stage" in fields else 0
                rows = assign_rows(body, bounds, len(fields), anchor)
                header_row = [clean_cell(cell["text"]) for cell in header_cells]
                span = (region[index]["top"], body[-1]["bottom"] if body else region[index]["bottom"])
                output.append(([header_row] + rows, span))
                index = max(end, index + consumed)
            else:
                index += 1
    return output


def _pdf_lines(model: PageModel, blocked: Sequence[tuple[float, float]]) -> list[TextLine]:
    lines: list[TextLine] = []
    for region in model.regions():
        previous = None
        for line in region:
            center = (line["top"] + line["bottom"]) / 2.0
            if any(top - 2 <= center <= bottom + 2 for top, bottom in blocked):
                previous = None
                continue
            text = clean_cell(line["text"])
            if not text:
                continue
            if previous is None:
                lines.append(TextLine(text, True))
            else:
                gap = line["top"] - previous["bottom"]
                height = max(6.0, previous["bottom"] - previous["top"])
                lines.append(TextLine(text, gap > 1.6 * height))
            previous = line
    return lines


def _extract_pdf(path: Path) -> tuple[list[TextLine], list[RawTable]]:
    text_lines: list[TextLine] = []
    raw_tables: list[RawTable] = []
    order = 0
    with pdfplumber.open(path) as pdf:
        for index, page in enumerate(pdf.pages):
            model = PageModel(page, index)
            found = ruled_tables(page)
            if not found:
                found = geometric_tables(model)
            blocked = []
            for rows, span in found:
                order += 1
                raw_tables.append(RawTable(headers=rows[0], rows=rows[1:], order=order, width=len(rows[0])))
                blocked.append(span)
            text_lines.extend(_pdf_lines(model, blocked))
    return text_lines, raw_tables


def _iter_docx_blocks(document: DocxDocument) -> Iterator[Any]:
    for child in document.element.body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            yield DocxParagraph(child, document)
        elif tag == "tbl":
            yield DocxTable(child, document)


def _extract_docx(path: Path) -> tuple[list[TextLine], list[RawTable]]:
    document = DocxDocument(path)
    text_lines: list[TextLine] = []
    raw_tables: list[RawTable] = []
    order = 0
    for block in _iter_docx_blocks(document):
        if isinstance(block, DocxParagraph):
            text = clean_cell(block.text)
            if text:
                heading = block.style is not None and "heading" in (block.style.name or "").lower()
                text_lines.append(TextLine(text, heading))
            continue
        rows = []
        for row in block.rows:
            seen = []
            cells = []
            for cell in row.cells:
                if cell._tc in seen:
                    continue
                seen.append(cell._tc)
                cells.append(clean_cell(cell.text))
            rows.append(cells)
        rows = [row for row in rows if any(row)]
        if rows:
            order += 1
            raw_tables.append(RawTable(headers=rows[0], rows=rows[1:], order=order, width=len(rows[0])))
    return text_lines, raw_tables


def _extract_markdown(path: Path) -> tuple[list[TextLine], list[RawTable]]:
    lines = Path(path).read_text(encoding="utf-8").splitlines()
    text_lines: list[TextLine] = []
    raw_tables: list[RawTable] = []
    buffer: list[str] = []
    order = 0

    def flush() -> None:
        nonlocal order, buffer
        if len(buffer) >= 2:
            rows = []
            for line in buffer:
                if re.fullmatch(r"\s*\|?[\s:\-|]+\|?\s*", line):
                    continue
                rows.append([clean_cell(cell) for cell in line.strip().strip("|").split("|")])
            rows = [row for row in rows if any(row)]
            if rows:
                order += 1
                raw_tables.append(RawTable(headers=rows[0], rows=rows[1:], order=order, width=len(rows[0])))
        else:
            for line in buffer:
                text = clean_cell(line)
                if text:
                    text_lines.append(TextLine(text, True))
        buffer = []

    for line in lines:
        if "|" in line:
            buffer.append(line)
            continue
        if buffer:
            flush()
        text = clean_cell(line.lstrip("#").strip())
        if text:
            text_lines.append(TextLine(text, line.startswith("#") or not line[:1].strip()))
    if buffer:
        flush()
    return text_lines, raw_tables


DOCUMENT_EXTRACTORS = {
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
    ".md": _extract_markdown,
    ".markdown": _extract_markdown,
    ".txt": _extract_markdown,
}


def _info_patterns() -> list[tuple[str, int, re.Pattern[str]]]:
    patterns = []
    for key, aliases in FIELD_ALIASES.items():
        for alias in sorted(aliases, key=len, reverse=True):
            body = r"\W+".join(re.escape(token) for token in alias.split())
            patterns.append((key, len(alias),
                             re.compile(r"^\W*%s\s*[:\uff1a\.\u2013\u2014\-]?\s*(.*)$" % body, re.IGNORECASE)))
    patterns.sort(key=lambda item: -item[1])
    return patterns


def _column_patterns() -> list[tuple[str, re.Pattern[str]]]:
    patterns = []
    for name, aliases in COLUMN_ALIASES.items():
        if name in {"index", "stage"}:
            continue
        body = "|".join(re.escape(alias) for alias in sorted(aliases, key=len, reverse=True))
        patterns.append((name, re.compile(r"^\W*(?:%s)\s*[:\uff1a\u2013\u2014\-]\s*(.+)$" % body, re.IGNORECASE)))
    return patterns


INFO_PATTERNS = _info_patterns()
COLUMN_PATTERNS = _column_patterns()


def match_field_label(line: str) -> Optional[tuple[str, str]]:
    stripped = clean_cell(line.lstrip("#").replace("**", "").replace("__", ""))
    if not stripped:
        return None
    for key, _, pattern in INFO_PATTERNS:
        found = pattern.match(stripped)
        if found:
            return key, clean_cell(found.group(1))
    return None


def parse_text_fields(lines: Sequence[TextLine]) -> dict[str, str]:
    results: dict[str, str] = {}
    active: Optional[str] = None
    buffer: list[str] = []

    def flush() -> None:
        nonlocal active, buffer
        if active and buffer:
            joined = " ".join(buffer).strip()
            if joined and not results.get(active):
                results[active] = joined
        active, buffer = None, []

    for line in lines:
        matched = match_field_label(line.text)
        if matched is not None:
            if results.get(matched[0]):
                flush()
                continue
            flush()
            active = matched[0]
            if matched[1]:
                buffer.append(matched[1])
            continue
        if active is None:
            continue
        if (line.break_before or SECTION_BREAK.match(line.text)
                or NUMBERED_HEADING.match(line.text)
                or any(pattern.match(line.text) for _, pattern in COLUMN_PATTERNS)):
            flush()
            continue
        buffer.append(line.text)
    flush()
    return results


def harvest_metadata_table(table: RawTable) -> dict[str, str]:
    harvested: dict[str, str] = {}
    for row in [table.headers] + table.rows:
        if len(row) < 2:
            continue
        matched = match_field_label(row[0])
        if matched is None:
            continue
        parts = [cell for cell in row[1:] if cell]
        value = " ".join([matched[1]] + parts).strip() if matched[1] else " ".join(parts).strip()
        if value:
            harvested.setdefault(matched[0], value)
    return harvested


def resolve_fields(fields: Sequence[Optional[str]]) -> list[Optional[str]]:
    present = {name for name in fields if name}
    if present & INVENTORY_FIELDS and not present & PROCESS_FIELDS:
        return ["equipment_inventory" if name == "equipment" else name for name in fields]
    return list(fields)


def classify_table(table: RawTable, fallback: Optional[list[Optional[str]]]
                   ) -> tuple[Optional[list[Optional[str]]], list[list[str]]]:
    fields = [match_column(cell) for cell in table.headers]
    if is_header_fields(fields, table.width):
        return resolve_fields(fields), table.rows
    if fallback and len(fallback) == table.width:
        anchor = fallback.index("stage") if "stage" in fallback else 0
        if anchor < len(table.headers) and looks_like_stage_name(table.headers[anchor]):
            return fallback, [table.headers] + table.rows
    return None, []


def table_records(fields: Sequence[Optional[str]], rows: Sequence[Sequence[str]]) -> list[dict[str, str]]:
    anchor = list(fields).index("stage") if "stage" in fields else 0
    grouped: list[list[str]] = []
    for row in rows:
        cells = [clean_cell(cell) for cell in row]
        while len(cells) < len(fields):
            cells.append("")
        if grouped and not cells[anchor]:
            for column, value in enumerate(cells):
                if value:
                    grouped[-1][column] = (grouped[-1][column] + " " + value).strip()
        else:
            grouped.append(cells)
    records = []
    for cells in grouped:
        name = cells[anchor]
        if not looks_like_stage_name(name):
            continue
        record = {"stage": name}
        for column, key in enumerate(fields):
            if not key or key in {"index", "stage"} or column >= len(cells):
                continue
            if cells[column]:
                record[key] = cells[column]
        if len(record) > 1:
            records.append(record)
    return records


def merge_stage_records(collected: list[dict[str, Any]], records: Sequence[dict[str, str]]) -> None:
    for record in records:
        key = normalize(record["stage"])
        target = None
        for existing in collected:
            if existing["_key"] == key or similar(existing["_key"], key) >= 0.88:
                target = existing
                break
        if target is None:
            target = {"_key": key, "stage": record["stage"]}
            collected.append(target)
        for name, value in record.items():
            if name != "stage" and value and not target.get(name):
                target[name] = value


def parse_stage_blocks(lines: Sequence[TextLine]) -> list[dict[str, str]]:
    records: list[dict[str, str]] = []
    current: Optional[dict[str, str]] = None
    active: Optional[str] = None
    for position, line in enumerate(lines):
        text = line.text
        if match_field_label(text) is not None and current is None:
            continue
        heading = NUMBERED_HEADING.match(text)
        candidate = clean_cell(heading.group(1)) if heading else text
        next_is_label = (position + 1 < len(lines)
                         and any(pattern.match(lines[position + 1].text) for _, pattern in COLUMN_PATTERNS))
        trailing = candidate.endswith((".", ",", ";", ":"))
        is_anchor = (looks_like_stage_name(candidate) and not trailing
                     and (heading is not None or (next_is_label and line.break_before)))
        if is_anchor and not any(pattern.match(text) for _, pattern in COLUMN_PATTERNS):
            current = {"stage": candidate}
            records.append(current)
            active = None
            continue
        if current is None:
            continue
        matched = None
        for name, pattern in COLUMN_PATTERNS:
            found = pattern.match(text)
            if found:
                matched = (name, clean_cell(found.group(1)))
                break
        if matched is not None:
            current.setdefault(matched[0], matched[1])
            active = matched[0]
            continue
        if active and not line.break_before and not SECTION_BREAK.match(text):
            current[active] = (current[active] + " " + text).strip()
    return [record for record in records if len(record) > 1]


def normalize_automation(value: str) -> str:
    text = clean_cell(value)
    if not text:
        return ""
    if text in {"\u2713", "\u2714", "\u2611", "\u25cf"}:
        return "Ya"
    if text in {"-", "\u2013", "\u2014", "\u2717", "\u2718", "\u2610"}:
        return "Tidak"
    tokens = normalize(text).split()
    if any(token in AUTOMATION_TRUE for token in tokens):
        return "Ya"
    if any(token in AUTOMATION_FALSE for token in tokens):
        return "Tidak"
    return text


def build_tables(stages: Sequence[dict[str, Any]], extras: Sequence[str]) -> list[ExtractedTable]:
    tables: list[ExtractedTable] = []
    for index, group in TABLE_GROUPS:
        columns = []
        for name in group:
            if name == "stage":
                columns.append(name)
                continue
            if name == "equipment_inventory" and not any(stage.get(name) for stage in stages):
                if any(stage.get("equipment") for stage in stages):
                    columns.append("equipment")
                continue
            if any(stage.get(name) for stage in stages):
                columns.append(name)
        if index == 1:
            columns = columns + list(extras)
        if len(columns) < 2:
            continue
        headers = [FIELD_HEADERS.get(name, name) for name in columns]
        rows = []
        for stage in stages:
            row = []
            for name in columns:
                value = stage.get(name, "")
                row.append(normalize_automation(value) if name == "automation" else value)
            rows.append(row)
        tables.append(ExtractedTable(index=index, headers=headers, rows=rows))
    return tables


def normalize_process_type(value: str) -> str:
    lowered = value.lower()
    if "paralel" in lowered or "parallel" in lowered:
        return "parallel"
    if "serial" in lowered or "sekuensial" in lowered or "sequential" in lowered:
        return "serial"
    return value.strip()


def normalize_worker_count(value: str) -> str:
    digits = re.search(r"\d+", value.replace(".", ""))
    return digits.group(0) if digits else value.strip()


def extract_document(path: str | Path) -> ExtractedDocument:
    resolved = Path(path)
    extractor = DOCUMENT_EXTRACTORS.get(resolved.suffix.lower())
    if extractor is None:
        raise UnsupportedDocumentError(
            f"Format dokumen tidak didukung: '{resolved.suffix}'. "
            f"Didukung: {sorted(set(DOCUMENT_EXTRACTORS) | WORKBOOK_SUFFIXES)}"
        )

    text_lines, raw_tables = extractor(resolved)

    stages: list[dict[str, Any]] = []
    harvested: dict[str, str] = {}
    extras: list[str] = []
    fallback: Optional[list[Optional[str]]] = None

    for table in sorted(raw_tables, key=lambda item: item.order):
        fields, rows = classify_table(table, fallback)
        if fields is None:
            harvested.update(harvest_metadata_table(table))
            continue
        fallback = fields
        for position, name in enumerate(fields):
            if name is None and position < len(table.headers) and table.headers[position]:
                label = table.headers[position]
                if label not in extras:
                    extras.append(label)
                fields[position] = label
        merge_stage_records(stages, table_records(fields, rows))

    if not stages:
        merge_stage_records(stages, parse_stage_blocks(text_lines))

    text_fields = parse_text_fields(text_lines)
    for key, value in harvested.items():
        if not text_fields.get(key):
            text_fields[key] = value

    if text_fields.get("process_type"):
        text_fields["process_type"] = normalize_process_type(text_fields["process_type"])
    if text_fields.get("worker_count"):
        text_fields["worker_count"] = normalize_worker_count(text_fields["worker_count"])

    for stage in stages:
        stage.pop("_key", None)

    return ExtractedDocument(
        source_name=resolved.name,
        text_fields=text_fields,
        tables=build_tables(stages, extras),
        stages=stages,
        raw_text="\n".join(line.text for line in text_lines),
    )


def build_agent_input(document: ExtractedDocument) -> str:
    fields = document.text_fields
    blocks = [
        f"Nama pabrik: {fields.get('factory_name', '')}",
        f"Jenis proses: {fields.get('process_type', '')}",
        f"Deskripsi pabrik: {fields.get('layout_description', '')}",
        f"Jumlah pekerja: {fields.get('worker_count', '')}",
        f"Urutan tahap proses: {' -> '.join(document.workflow_sequence())}",
    ]
    for table in document.tables:
        label = TABLE_LABELS.get(table.index, f"TABEL {table.index}")
        blocks.append(f"\n{label}\n{table.to_markdown()}")
    return "\n".join(blocks)